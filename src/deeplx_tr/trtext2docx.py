"""Generate docx for source_text, trtext, alt_text."""
# pylint: disable=too-many-locals, broad-exception-caught, too-many-branches, too-many-statements,

import os
from itertools import zip_longest
from pathlib import Path
from secrets import token_hex
from typing import List

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX

# from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from loadtext import loadtext
from loguru import logger


def trtext2docx(
    source_text: None | List[str] | str = None,
    trtext: None | List[str] | str = None,
    alt_text: None | List[str] | str = None,
    outfile: str | None = None,
    openfile: bool = True,
):
    """
    Generate docx for source_text, trtext, alt_text.

    Args:
    ----
    source_text: source text
    trtext: translated text
    alt_text: alternative translated text
    outfile: file path if True,
    openfile: try to open the saved file if True, default True

    Returns:
    -------
    docx.Document and saved file if file path supplied

    """
    if not source_text:
        source_text = []

    if not trtext:
        trtext = []

    if not alt_text:
        alt_text = []

    if isinstance(source_text, str):
        source_text = source_text.splitlines()

    if isinstance(trtext, str):
        trtext = trtext.splitlines()

    if isinstance(alt_text, str):
        alt_text = alt_text.splitlines()

    templ_dual = Path(__file__).parent / "templ_dual.docx"
    if templ_dual.exists():
        document = Document(templ_dual.as_posix())
        logger.info(f"Using {templ_dual=}")
    else:
        document = Document()

    # Normal style: built-in
    document.styles["Normal"].font.name = "宋体"
    document.styles["Normal"].font.highlight_color = WD_COLOR_INDEX.YELLOW
    # document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    # document.styles['Normal'].font.size = Pt(7)
    # document.styles['Normal'].font.size = Pt(12)
    document.styles["Normal"].paragraph_format.line_spacing = Pt(0)

    # TrtextStyle
    trtext_style = document.styles.add_style("TrtextStyle", WD_STYLE_TYPE.PARAGRAPH)

    trtext_font = trtext_style.font
    trtext_font.name = "宋体"
    trtext_font.size = Pt(10)

    # Create Alttext style
    alttext_style = document.styles.add_style("AlttextStyle", WD_STYLE_TYPE.PARAGRAPH)
    alttext_font = alttext_style.font
    alttext_style.name = "Times New Roman"
    alttext_font.size = Pt(10)
    alttext_font.highlight_color = WD_COLOR_INDEX.WHITE  # YELLOW WHITE GRAY_25
    alttext_font.color.rgb = RGBColor(0xFF, 0x0, 0xE0)
    paragraph = document.add_paragraph("", style="Normal")

    for col0, col1, col2 in zip_longest(source_text, trtext, alt_text):
        if col0:
            # paragraph = document.add_paragraph("", style="Normal")
            paragraph = document.add_paragraph(col0, style="Normal")
        if col1:
            paragraph.paragraph_format.space_after = Pt(12)
            paragraph = document.add_paragraph(col1, style="TrtextStyle")
            # paragraph.paragraph_format.space_after = Pt(10.1)
        if col2:
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph = document.add_paragraph(col2, style="AlttextStyle")

    if outfile:
        try:
            document.save(outfile)
            logger.info(f"Saved to {outfile=}")
        except Exception as exc:
            logger.warning(f" Cant save to {outfile=}: {exc}")
        if openfile:
            try:
                os.startfile(outfile)
            except Exception as exc:
                logger.info(f"Cant open {outfile=}: {exc}")

    return document


def main():  # pylint: disable=missing-function-docstring
    texts = loadtext(r"C:\syncthing\00xfer\2021it\2024-05-30.txt")

    ofile = f"temp-{token_hex(3)}.docx"
    trtext2docx(texts, texts, texts, ofile)


if __name__ == "__main__":
    main()
